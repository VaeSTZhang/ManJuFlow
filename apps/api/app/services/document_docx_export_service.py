from io import BytesIO
from pathlib import PurePosixPath
import re
from typing import Any

from docx import Document

from app.schemas.document import DocumentExportInput
from app.services.document_usage_ledger import record_document_export_usage
from app.services.ownership_service import record_export_document_ownership


DEFAULT_DOCX_BASENAME = "dramora-document"


def _safe_filename_part(value: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", value.strip(), flags=re.UNICODE)
    cleaned = cleaned.strip(".-")
    return cleaned or DEFAULT_DOCX_BASENAME


def _strip_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[0] if "." in filename else filename


def sanitize_docx_filename(filename: str | None, project_title: str) -> str:
    if filename:
        normalized_filename = filename.replace("\\", "/").strip()
        safe_name = PurePosixPath(normalized_filename).name.strip()
        base_name = _strip_extension(safe_name)
    else:
        base_name = project_title

    return f"{_safe_filename_part(base_name)}.docx"


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _add_optional_paragraph(document: Document, label: str, value: Any) -> None:
    if value is None:
        return

    text = str(value).strip()
    if not text:
        return

    document.add_paragraph(f"{label}：{text}")


def _add_text_block(document: Document, content_text: str) -> None:
    for line in content_text.splitlines():
        text = line.strip()
        if text:
            document.add_paragraph(text)


def _add_text_list(document: Document, title: str, items: list[Any]) -> None:
    visible_items = [str(item).strip() for item in items if str(item).strip()]
    if not visible_items:
        return

    document.add_heading(title, level=2)
    for item in visible_items:
        document.add_paragraph(item, style="List Bullet")


def _add_characters(document: Document, payload: dict[str, Any]) -> None:
    characters = _as_list(payload.get("characters"))
    if not characters:
        return

    document.add_heading("主要人物", level=2)
    for character_value in characters:
        character = _as_dict(character_value)
        name = str(character.get("name") or "未命名角色")
        role = str(character.get("role") or "")
        title = f"{name}（{role}）" if role else name
        document.add_paragraph(title, style="List Bullet")
        _add_optional_paragraph(document, "年龄", character.get("age"))
        _add_optional_paragraph(document, "性格", character.get("personality"))
        _add_optional_paragraph(document, "人物弧光", character.get("arc"))


def _add_adaptation_notes(document: Document, payload: dict[str, Any]) -> None:
    notes = _as_dict(payload.get("adaptation_notes"))
    if not notes:
        return

    document.add_heading("改编策略", level=2)
    _add_optional_paragraph(document, "策略", notes.get("adaptation_strategy"))
    _add_text_list(document, "保留元素", _as_list(notes.get("preserved_elements")))
    _add_text_list(document, "调整内容", _as_list(notes.get("changed_elements")))
    _add_text_list(document, "短剧钩子 / 爆点", _as_list(notes.get("short_drama_hooks")))
    _add_text_list(document, "备注", _as_list(notes.get("risk_notes")))


def _add_dialogues(document: Document, dialogues: list[Any]) -> None:
    visible_dialogues = [_as_dict(dialogue) for dialogue in dialogues]
    if not visible_dialogues:
        return

    document.add_paragraph("对白：")
    for dialogue in visible_dialogues:
        character = str(dialogue.get("character") or "角色")
        line = str(dialogue.get("line") or "").strip()
        if line:
            document.add_paragraph(f"{character}：{line}", style="List Bullet")


def _add_scenes(document: Document, episode: dict[str, Any]) -> None:
    scenes = _as_list(episode.get("scenes"))
    if not scenes:
        return

    for scene_value in scenes:
        scene = _as_dict(scene_value)
        scene_number = scene.get("scene_number")
        scene_title = f"第 {scene_number} 场" if scene_number is not None else "场景"
        location = str(scene.get("location") or "").strip()
        time = str(scene.get("time") or "").strip()
        suffix = "｜".join([item for item in [location, time] if item])
        document.add_heading(f"{scene_title}{'｜' + suffix if suffix else ''}", level=3)
        _add_optional_paragraph(document, "场景", scene.get("description"))
        _add_dialogues(document, _as_list(scene.get("dialogues")))
        _add_optional_paragraph(document, "画面备注", scene.get("visual_notes"))
        _add_optional_paragraph(document, "情绪备注", scene.get("emotion_curve"))


def _add_episodes(document: Document, payload: dict[str, Any]) -> None:
    episodes = _as_list(payload.get("episodes"))
    if not episodes:
        return

    document.add_heading("分集内容", level=2)
    for episode_value in episodes:
        episode = _as_dict(episode_value)
        episode_number = episode.get("episode_number")
        title = str(episode.get("title") or "未命名分集")
        heading = f"第 {episode_number} 集：{title}" if episode_number is not None else title
        document.add_heading(heading, level=3)
        _add_optional_paragraph(document, "概要", episode.get("summary"))
        _add_optional_paragraph(document, "钩子", episode.get("hook"))
        _add_scenes(document, episode)


def build_short_drama_docx_document(input_data: DocumentExportInput) -> Document:
    document = Document()
    payload = _as_dict(input_data.structured_payload)

    title = str(payload.get("project_title") or input_data.project_title)
    document.add_heading(title, level=1)

    if payload:
        _add_optional_paragraph(document, "故事梗概", payload.get("logline"))
        _add_optional_paragraph(document, "世界观 / 故事背景", payload.get("world_setting"))
        _add_characters(document, payload)
        _add_adaptation_notes(document, payload)
        _add_episodes(document, payload)
    elif input_data.content_text:
        _add_text_block(document, input_data.content_text)
    else:
        document.add_paragraph("无可导出内容。")

    return document


def build_docx_export_bytes(input_data: DocumentExportInput) -> bytes:
    if input_data.export_format != "docx":
        record_document_export_usage(
            input_data=input_data,
            document_operation="export_docx",
            status="failed",
            error_code="document_export_docx_failed",
            error_message_safe="DOCX export requires export_format='docx'.",
        )
        raise ValueError("DOCX export requires export_format='docx'.")

    try:
        document = build_short_drama_docx_document(input_data)
        buffer = BytesIO()
        document.save(buffer)
        docx_bytes = buffer.getvalue()
    except Exception as exc:
        record_document_export_usage(
            input_data=input_data,
            document_operation="export_docx",
            status="failed",
            error_code="document_export_docx_failed",
            error_message_safe="DOCX export failed.",
        )
        raise ValueError("DOCX export failed.") from exc

    record_document_export_usage(
        input_data=input_data,
        document_operation="export_docx",
        file_size_bytes=len(docx_bytes),
    )
    record_export_document_ownership(
        context_options=input_data.context_options,
        project_title=input_data.project_title,
        source_mode=_structured_source_mode(input_data),
        document_type=input_data.document_type,
        source_stage=input_data.source_stage,
        export_format="docx",
        filename_safe=sanitize_docx_filename(input_data.filename, input_data.project_title),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size_bytes=len(docx_bytes),
        character_count=None,
        episode_count=_structured_count(input_data, "episodes"),
        characters_count=_structured_count(input_data, "characters"),
    )
    return docx_bytes


def _structured_source_mode(input_data: DocumentExportInput) -> str | None:
    if input_data.structured_payload is None:
        return None
    source_mode = input_data.structured_payload.get("source_mode")
    return str(source_mode) if source_mode else None


def _structured_count(input_data: DocumentExportInput, key: str) -> int | None:
    if input_data.structured_payload is None:
        return None
    value = input_data.structured_payload.get(key)
    if isinstance(value, list):
        return len(value)
    return None
